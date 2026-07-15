from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = "output/Melann_Project_System_Implementation_Guide.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CAUTION = "FFF2CC"
GREEN = "E2F0D9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("IMPLEMENTATION GUIDE")
    set_run_font(r, size=10, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Melann Lending Past Due and Report Monitoring")
    set_run_font(r, size=24, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Printable step-by-step guide for branch implementation and JCASH migration")
    set_run_font(r, size=12, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, [1.55, 4.95])
    rows = [
        ("System", "React/Vite frontend, Express bridge server, PostgreSQL database"),
        ("Default URLs", "Frontend http://localhost:3000, API http://localhost:5000/api"),
        ("Primary Branch Use", "Naval Branch, Ormoc Branch, or All Branches depending on user role"),
        ("Prepared For", "Implementation in a separate branch/workspace before live use"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    add_callout(
        doc,
        "Important",
        "Treat this as a runbook for tomorrow's implementation. Do the migration in the new branch/workspace first, confirm backup and database target, then migrate JCASH data only after the smoke checks pass.",
        CAUTION,
    )


def add_callout(doc, title, body, fill):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{title}: ")
    r.bold = True
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run("[ ] ").bold = True
        p.add_run(item)


def add_steps(doc, steps):
    for title, detail in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(title)
        r.bold = True
        if detail:
            p.add_run(f" - {detail}")


def add_command(doc, command):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(command)
    set_run_font(r, name="Consolas", size=9.5, color=RGBColor(31, 31, 31))


def add_phase_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [0.7, 1.35, 3.15, 1.3])
    headers = ["Done", "Phase", "Action / Evidence", "Owner"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.add_run(header).bold = True
    for phase, action, owner in rows:
        cells = table.add_row().cells
        cells[0].text = "[ ]"
        cells[1].text = phase
        cells[2].text = action
        cells[3].text = owner
        for cell in cells:
            set_cell_margins(cell)


def add_doc_footer(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Melann Project System Implementation Guide")
        set_run_font(run, size=9, color=MUTED)


def build():
    doc = Document()
    set_doc_defaults(doc)
    add_title_block(doc)

    doc.add_heading("1. Objective", level=1)
    doc.add_paragraph(
        "Goal nito: ma-implement ang Project System sa ibang branch/workspace nang hindi nagagalaw ang kasalukuyang working copy, ma-setup ang database, ma-run ang app, ma-migrate ang JCASH accounts, at ma-verify na tama ang data bago gamitin ng branch users."
    )

    doc.add_heading("2. Before Implementation Day", level=1)
    add_checklist(doc, [
        "Siguraduhin na may hiwalay na Git branch o hiwalay na folder/workspace para sa implementation.",
        "Kumpirmahin ang target database connection string para sa branch. Huwag gamitin ang production database kung dry-run pa lang.",
        "Kumuha ng database backup mula sa PostgreSQL provider or existing branch database tools.",
        "Kumpirmahin na accessible ang JCASH MDB source path: \\\\SERVERPC\\LendingV2Melan\\db\\jcashdb.mdb.",
        "Kumpirmahin kung kailangan ng JCASHDB_PASSWORD at ilagay lang ito sa .env.local, hindi sa code.",
        "Ihanda ang admin/super-admin login na gagamitin sa migration at user approval.",
    ])

    doc.add_heading("3. Implementation Sequence Overview", level=1)
    add_phase_table(doc, [
        ("Code", "Create/switch to the implementation branch and confirm the repo is clean enough to work on.", "Implementer"),
        ("Install", "Run npm install if node_modules is missing or dependencies changed.", "Implementer"),
        ("Env", "Create or update .env.local with DATABASE_URL, PORT, optional VITE_GEMINI_API_KEY, and JCASH settings.", "Implementer"),
        ("DB", "Run schema setup against the target database and confirm default admin/user tables are present.", "Implementer"),
        ("Run", "Start backend and frontend; verify http://localhost:3000 and http://localhost:5000/api/loans.", "Implementer"),
        ("Backup", "Take/confirm backup immediately before migration/import actions.", "Implementer"),
        ("JCash", "Use JCASH Migration screen to scan, review, select, and migrate accounts.", "Super Admin"),
        ("Verify", "Check migrated counts, sample accounts, payments, balances, reports, and branch user access.", "Implementer"),
        ("Handoff", "Approve users, share launcher/startup process, and record any unresolved issue.", "Admin"),
    ])

    doc.add_heading("4. Branch and Workspace Setup", level=1)
    add_steps(doc, [
        ("Open the project folder", "Use the MelannPastDueReportMonitoring workspace."),
        ("Create or switch to the implementation branch", "Use a branch name that clearly identifies the target branch/location and date."),
        ("Check working tree status", "Confirm which files are already modified before making implementation changes."),
    ])
    add_command(doc, "git status")
    add_command(doc, "git switch -c implementation/<branch-name>-<date>")
    add_callout(doc, "Note", "If the branch already exists, use git switch <branch-name>. Do not reset or discard existing work unless the owner explicitly approves.", LIGHT_GRAY)

    doc.add_heading("5. Install and Configure", level=1)
    add_steps(doc, [
        ("Install dependencies", "Run this only when node_modules is missing or package-lock/package.json changed."),
        ("Create .env.local", "Use the target branch database URL and safe local server port."),
        ("Do not commit secrets", ".env.local must stay local and should not be committed."),
    ])
    add_command(doc, "npm install")
    add_command(doc, "DATABASE_URL=postgresql://USER:PASSWORD@HOST:PORT/DATABASE\nPORT=5000\nVITE_GEMINI_API_KEY=optional_gemini_key\nJCASHDB_PATH=\\\\SERVERPC\\LendingV2Melan\\db\\jcashdb.mdb\nJCASHDB_PASSWORD=optional_if_required\nJCASHDB_BRANCH=Ormoc Branch")

    doc.add_heading("6. Database Setup", level=1)
    doc.add_paragraph("The database schema is defined in schema.sql. The setup script reads DATABASE_URL from .env.local and applies the schema.")
    add_command(doc, "node scripts/setup_db.js")
    add_checklist(doc, [
        "Confirm successful schema setup message.",
        "Confirm the backend can connect to PostgreSQL.",
        "Confirm users, collectors, loans, payments, remarks, demand_letters, activity_logs, visit_logs, deleted_loans, and migration_batches tables exist.",
        "If this is an existing database, confirm no unintended wipe/import action has been run.",
    ])

    doc.add_heading("7. Start the System", level=1)
    doc.add_paragraph("Use two terminals. On Windows PowerShell, prefer npm.cmd commands.")
    add_command(doc, "npm.cmd run server")
    add_command(doc, "npm.cmd run dev")
    doc.add_paragraph("Expected backend output includes database connection success and Bridge Server running on port 5000. Open the frontend at http://localhost:3000.")
    add_callout(doc, "Shortcut", "Start_Melann_System.bat can launch backend, frontend, wait, then open Chrome. Use it only after .env.local and database setup are correct.", LIGHT_GRAY)

    doc.add_heading("8. Pre-Migration Safety Checks", level=1)
    add_checklist(doc, [
        "Take or confirm database backup before any import, wipe, cleanup, or migration.",
        "Login as admin or a user with data maintenance access.",
        "Open Dashboard and Loan Grid; confirm existing data is expected for the target branch.",
        "Confirm API works by opening http://localhost:5000/api/loans.",
        "Run frontend build to catch compile issues before migration.",
    ])
    add_command(doc, "npm.cmd run build")

    doc.add_heading("9. JCASH Migration - Preferred In-App Flow", level=1)
    doc.add_paragraph("Use the JCASH Migration screen instead of one-off scripts when possible. The backend creates a read-only snapshot copy of the MDB first, scans by maturity date, stores pending results in migration_batches, then imports selected accounts and good payments.")
    add_steps(doc, [
        ("Open JCASH Migration", "Use the sidebar item under Administration/Data Maintenance."),
        ("Set Maturity Date From and To", "Example active portfolio default in the app is 2016-01-01 to 2026-03-31."),
        ("Click Scan Database", "The system reads from the configured JCASHDB_PATH and creates a pending batch."),
        ("Review detected accounts", "Check borrower, code, due date, principal, total loan, collector, branch, address, contact number, balance, status, and good payment count."),
        ("Edit incorrect client details if needed", "Use the edit action before migration; edited rows are marked in the pending batch."),
        ("Exclude accounts that should not migrate", "Excluded accounts move to Recycle Bin/deleted_loans with reason Excluded from JCASH Migration."),
        ("Select accounts", "Use Select All, manual checkboxes, or Select 0 Pay/NMSR for zero-payment or NMSR accounts."),
        ("Migrate selected clients", "The system upserts loans, replaces prior JCASH-migrated payments for those loans, imports good payments, and updates the remaining pending count."),
        ("Record the result", "Write down importedCount, paymentCount, remainingCount, date/time, and migratedBy."),
    ])
    add_callout(doc, "Payment behavior", "JCASH migrated payments use remarks 'Migrated from jcashdb.mdb'. During migration for a loan, existing payments with that same remark are deleted and reinserted from the selected batch, while non-JCASH/manual payments are preserved.", CAUTION)

    doc.add_heading("10. Optional Script-Based Migration / Repair Tools", level=1)
    doc.add_paragraph("Use these only when the in-app migration is not enough or a specific repair is required. Read the script before running. Prefer dry-run mode where available.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [2.2, 2.55, 1.75])
    for idx, header in enumerate(["Script", "Purpose", "Safety Rule"]):
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
        table.cell(0, idx).paragraphs[0].add_run(header).bold = True
    scripts = [
        ("migrate.cjs", "Legacy CSV-based loans/payments migration from loans.csv and payments.csv.", "Review DB string first."),
        ("migrate_specific_payments.cjs", "Migrates payments only for codes 1064, 1254, 957, 1423.", "Review mappings first."),
        ("scripts/fix_zero_payment_paid_jcash_loans.cjs", "Finds or fixes Ormoc zero-payment positive-balance JCASH loans.", "Run without --apply first."),
        ("scripts/fix_jcash_total_loan_mismatches.cjs", "Corrects target JCASH total loan mismatches using output/jcash_compare_data.json.", "Run without --apply first."),
        ("scripts/fix_3544_missing_jcash_payments.cjs", "Adds known missing JCASH payments for client code 3544.", "Run without --apply first."),
        ("verify_db.cjs / verify_good_status.cjs", "Database and status verification helpers.", "Read output carefully."),
    ]
    for script, purpose, rule in scripts:
        cells = table.add_row().cells
        cells[0].text = script
        cells[1].text = purpose
        cells[2].text = rule
        for cell in cells:
            set_cell_margins(cell)
    add_callout(doc, "Warning", "Some older root-level migration scripts contain hardcoded local connection strings. Update them or prefer .env.local-based scripts before running against a real target database.", CAUTION)

    doc.add_heading("11. Post-Migration Verification", level=1)
    add_checklist(doc, [
        "Refresh the app and confirm migrated accounts appear in the correct branch.",
        "Sample at least 10 migrated accounts: check code, borrower name, release date, due date, principal, total loan, running balance, status, collector, address, and contact number.",
        "For accounts with payments, open profile/payment history and compare count and final balance against the JCASH scan output.",
        "For zero-payment or NMSR accounts, confirm balance was not accidentally set to zero unless actually paid.",
        "Open Dashboard, Loan Grid, Collection Sheet, Daily Collection Report, Reports, Aging Report, Client Update, and Backup/Restore.",
        "Check browser console and backend terminal for errors.",
        "Run build again after any last-minute fixes.",
    ])
    add_command(doc, "npm.cmd run build")

    doc.add_heading("12. User Access and Branch Handoff", level=1)
    add_steps(doc, [
        ("Login as admin", "The default seeded user is admin when users table starts empty."),
        ("Approve or create branch users", "Use User Management and assign correct role and branch."),
        ("Confirm branch scoping", "Naval users see Naval Branch, Ormoc users see Ormoc Branch, super-admin can select All Branches."),
        ("Show startup process", "Demonstrate two-terminal startup or Start_Melann_System.bat."),
        ("Export backup", "Use Backup & Restore after migration as a post-migration snapshot."),
    ])

    doc.add_heading("13. Final Go/No-Go Checklist", level=1)
    add_checklist(doc, [
        "Correct Git branch/workspace is active.",
        "Correct target DATABASE_URL is configured.",
        "Database backup exists before migration.",
        "Backend starts without database errors.",
        "Frontend opens at http://localhost:3000.",
        "Build passes or any known build issue is documented.",
        "JCASH scan completed for the intended maturity date range.",
        "Selected accounts migrated with recorded imported/payment/remaining counts.",
        "Sample account and payment verification passed.",
        "Reports and branch user access passed smoke testing.",
        "Post-migration backup/export completed.",
        "Outstanding issues and excluded accounts are documented.",
    ])

    doc.add_heading("14. Quick Troubleshooting", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [1.65, 2.45, 2.4])
    for idx, header in enumerate(["Issue", "Likely Cause", "Action"]):
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
        table.cell(0, idx).paragraphs[0].add_run(header).bold = True
    rows = [
        ("Frontend has no data", "Backend/database unavailable", "Check npm.cmd run server, .env.local, and http://localhost:5000/api/loans."),
        ("Login fails", "User pending/deactivated or not seeded", "Check User Management or users table; admin seeds only when users table is empty."),
        ("JCASH scan fails", "MDB path, permission, ACE OLEDB, password, or temp folder issue", "Verify JCASHDB_PATH, network share access, JCASHDB_PASSWORD, and C:\\tmp availability."),
        ("Duplicate payment conflict", "Same loan/date or OR number collision", "Use in-app migration where possible; review one-off payment scripts before apply."),
        ("Wrong branch after migration", "JCASHDB_BRANCH or selected target branch is wrong", "Correct environment/configuration before re-scanning/migrating."),
        ("PowerShell blocks npm run", "Script policy", "Use npm.cmd run build/dev/server."),
    ]
    for issue, cause, action in rows:
        cells = table.add_row().cells
        cells[0].text = issue
        cells[1].text = cause
        cells[2].text = action
        for cell in cells:
            set_cell_margins(cell)

    add_doc_footer(doc)
    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
